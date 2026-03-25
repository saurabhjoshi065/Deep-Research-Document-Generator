"""Word document compiler."""

from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt
from src.state import SectionDraft

class DocxCompiler:
    """Compiler for DOCX output format."""

    def compile(self, sections: List[SectionDraft], title: str, output_path: Path, metadata: Optional[Dict] = None) -> Path:
        """Compile sections into a Word document."""
        doc = Document()
        
        # Title
        title_obj = doc.add_heading(title, 0)
        
        # TOC Placeholder (Manual in Word usually, but we list it)
        doc.add_heading("Table of Contents", level=1)
        for i, s in enumerate(sections, 1):
            doc.add_paragraph(f"{i}. {s.section_title}", style='List Bullet')
        
        doc.add_page_break()

        # Sections
        for s in sections:
            doc.add_heading(s.section_title, level=1)
            doc.add_paragraph(s.content)
            doc.add_page_break()

        # Sources
        doc.add_heading("Sources", level=1)
        all_sources = set()
        for s in sections: all_sources.update(s.sources_cited)
        for src in sorted(all_sources):
            doc.add_paragraph(src, style='List Bullet')

        doc.save(str(output_path))
        return output_path
